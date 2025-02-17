import csv
import datetime
import io
import json
import uuid

import pypdfium2
from docx import Document
from markdown_it import MarkdownIt
from openpyxl import load_workbook

from base_models import DocumentData
from settings import MAX_ROW

HANDLERS = {}


def register_extractor(*extensions):
    def decorator(func):
        for ext in extensions:
            HANDLERS[ext] = func
        return func

    return decorator


@register_extractor('pdf')
def extract_pdf(content, filename):
    CURRENT_ISO_TIME = datetime.datetime.now(datetime.UTC).isoformat(timespec='milliseconds')
    unique_id = str(uuid.uuid4())

    # Open the PDF document
    pdf = pypdfium2.PdfDocument(content)

    # Iterate through all pages in the PDF
    for page_num in range(len(pdf)):
        page = pdf[page_num]

        # Extract text from the page
        textpage = page.get_textpage()
        text = textpage.get_text_range()

        # Split the text into lines (pypdfium2 separates lines correctly)
        lines = text.split('\n')

        # Yield each line for further processing
        for line_number, line in enumerate(lines, start=1):
            yield (DocumentData(id=unique_id, document_type="pdf", document_name=filename, raw_content=line,
                                line_no=line_number,
                                page_no=page_num + 1, uploaded_on=CURRENT_ISO_TIME))


@register_extractor('xls', 'xlsm')
def extract_xlsm(content, filename, start_row=1, end_row=MAX_ROW):
    CURRENT_ISO_TIME = datetime.datetime.now(datetime.UTC).isoformat(timespec='milliseconds')
    unique_id = str(uuid.uuid4())
    # Load the XLSM workbook
    wb = load_workbook(io.BytesIO(content), data_only=True)
    sheet = wb.active  # Get the active sheet

    # Find the number of columns by inspecting the first row with data
    max_column = sheet.max_column  # Get the maximum column with data in the sheet

    # Iterate through the specified range of rows
    for row_number in range(start_row, end_row + 1):
        # Iterate through the cells in the row up to the `max_column`
        line = " ".join(str(sheet.cell(row=row_number, column=col).value)
                        if sheet.cell(row=row_number, column=col).value is not None else ""
                        for col in range(1, max_column + 1))

        # Yield the row number and the line content
        yield DocumentData(id=unique_id, document_type="xlsm", document_name=filename, raw_content=line,
                           line_no=row_number,
                           uploaded_on=CURRENT_ISO_TIME)


@register_extractor('csv')
def extract_csv(content, filename, start_row=1, end_row=MAX_ROW):
    CURRENT_ISO_TIME = datetime.datetime.now(datetime.UTC).isoformat(timespec='milliseconds')
    unique_id = str(uuid.uuid4())
    text_stream = io.StringIO(content.decode('utf-8'))
    reader = csv.reader(text_stream)
    # Iterate through rows, and stop when we reach end_row
    for row_number, row in enumerate(reader, start=start_row):
        if row_number > end_row:
            break

        # Get the number of columns in this row dynamically
        line = ", ".join(str(cell) if cell is not None else "" for cell in row)

        # Yield the row number and the line content
        yield DocumentData(id=unique_id, document_type="csv", document_name=filename, raw_content=line,
                           line_no=row_number,
                           uploaded_on=CURRENT_ISO_TIME)


@register_extractor('json')
def extract_json(content, filename):
    CURRENT_ISO_TIME = datetime.datetime.now(datetime.UTC).isoformat(timespec='milliseconds')
    unique_id = str(uuid.uuid4())
    json_string = content.decode('utf-8')
    return [
        DocumentData(id=unique_id, document_type="json", document_name=filename, raw_content=json_string, line_no=0,
                     uploaded_on=CURRENT_ISO_TIME)]


@register_extractor('docx')
def extract_docx(content, filename, start_line=0, end_line=MAX_ROW):
    CURRENT_ISO_TIME = datetime.datetime.now(datetime.UTC).isoformat(timespec='milliseconds')
    unique_id = str(uuid.uuid4())
    # Load the document
    doc = Document(io.BytesIO(content))

    # Iterate through paragraphs and return the line number and content
    for line_number, para in enumerate(doc.paragraphs[start_line:end_line], start=start_line + 1):
        # Yield the line number and the paragraph text
        yield DocumentData(id=unique_id, document_type='docx', document_name=filename, raw_content=para.text,
                           line_no=line_number,
                           uploaded_on=CURRENT_ISO_TIME)


@register_extractor('md')
def extract_markdown(content, filename):
    CURRENT_ISO_TIME = datetime.datetime.now(datetime.UTC).isoformat(timespec='milliseconds')
    unique_id = str(uuid.uuid4())
    md = MarkdownIt()
    decoded_content = content.decode('utf-8')

    # Split the content into lines
    lines = decoded_content.splitlines()

    # Process each line
    for line_number, line in enumerate(lines, start=1):
        # Remove leading/trailing whitespace
        line = line.strip()

        # Render the line to HTML using MarkdownIt
        rendered_html = md.render(line)

        # Yield the result (assuming you're streaming the data back)
        yield DocumentData(
            id=unique_id,
            document_type="md",
            document_name=filename,
            raw_content=line,
            line_no=line_number,
            content=rendered_html,
            uploaded_on=CURRENT_ISO_TIME
        )


@register_extractor('txt')
def extract_text(content, filename):
    CURRENT_ISO_TIME = datetime.datetime.now(datetime.UTC).isoformat(timespec='milliseconds')
    unique_id = str(uuid.uuid4())

    lines = content.decode('utf-8').splitlines()
    result = [
        DocumentData(
            id=unique_id,
            document_type="txt",
            document_name=filename,
            raw_content=line.strip(),
            line_no=line_number + 1,
            uploaded_on=CURRENT_ISO_TIME
        )
        for line_number, line in enumerate(lines)
    ]

    return result

# async def main():
# pdf
# pdf_path = 'data/sample.pdf'
# for page_num, line_num, line in extract_pdf(pdf_path):
#     print(f"Page {page_num}, Line {line_num}: {line}")

# xlsm
# xlsx_path = '/Users/254428/PycharmProjects/igloo/data/ota.xlsm'
# for row_num, line in extract_xlsm(xlsx_path):
#     print(f"Row {row_num}: {line.expandtabs()}")

# csv
# csv_path = 'data/cui.csv'
# for row_num, line in extract_csv(csv_path):
#     print(f"Row {row_num}: {line}")

# json
# print(extract_json("data/metadata.json"))

# docx
# docx_path = 'data/sample.docx'
# for line_num, line_text in extract_docx(docx_path):
#     print(f"Line {line_num}: {line_text}")


# markdown
# md_path = 'data/sample.md'
# for line_num, line_text,line_html in extract_markdown(md_path):
#     print(f"Line {line_num}: {line_text} {line_html}")

# text
# file_path = '../data/sample.txt'
# for line_number, line in extract_text(file_path):
#     print(f"Line {line_number}: {line}")


# asyncio.run(main())
